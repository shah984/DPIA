"""
generator.py — Populate dpia_template.docx with extracted field values and save output.

Opens a fresh copy of the template (never modifies the original), replaces:
  - Text-input SDT content controls (by sequential positional index)
  - Date-picker SDT content controls (by sequential positional index)
  - Checkbox SDT glyphs ☐ → ☒ (by sequential positional index among all ☐ SDTs)
  - Plain table cells (cover table)

Saves to: outputs/DPIA_<stem>_<YYYY-MM-DD>.docx
"""

from __future__ import annotations

import logging
from datetime import date
from pathlib import Path

from docx import Document
from lxml import etree

from .dpia_mapper import MappingPlan

logger = logging.getLogger(__name__)

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_TEMPLATE = Path(__file__).resolve().parent.parent / "dpia_docs" / "dpia_template.docx"
_OUTPUTS  = Path(__file__).resolve().parent.parent / "outputs"


# ── SDT helpers ───────────────────────────────────────────────────────────────

def _collect_text_sdts(doc: Document) -> list:
    """Return body-level SDTs whose placeholder is 'Click or tap here to enter text.'"""
    result = []
    for sdt in doc.element.body.iter(f"{{{W}}}sdt"):
        for t in sdt.iter(f"{{{W}}}t"):
            if t.text and "Click or tap here to enter text" in t.text:
                result.append(sdt)
                break
    return result


def _collect_date_sdts(doc: Document) -> list:
    """Return body-level SDTs whose placeholder is 'Click or tap to enter a date.'"""
    result = []
    for sdt in doc.element.body.iter(f"{{{W}}}sdt"):
        for t in sdt.iter(f"{{{W}}}t"):
            if t.text and "Click or tap to enter a date" in t.text:
                result.append(sdt)
                break
    return result


def _collect_checkbox_sdts(doc: Document) -> list:
    """Return all SDTs whose sdtContent is solely the ☐ glyph (U+2610)."""
    BALLOT = "\u2610"
    result = []
    for sdt in doc.element.body.iter(f"{{{W}}}sdt"):
        content = sdt.find(f"{{{W}}}sdtContent")
        if content is None:
            continue
        texts = [t.text for t in content.iter(f"{{{W}}}t") if t.text]
        if "".join(texts).strip() == BALLOT:
            result.append(sdt)
    return result


def _tick_checkbox(sdt) -> None:
    """Change the ☐ glyph in an SDT's sdtContent to ☒ (U+2612)."""
    BALLOT_EMPTY  = "\u2610"
    BALLOT_TICKED = "\u2612"
    content = sdt.find(f"{{{W}}}sdtContent")
    if content is None:
        return
    for t in content.iter(f"{{{W}}}t"):
        if t.text and BALLOT_EMPTY in t.text:
            t.text = t.text.replace(BALLOT_EMPTY, BALLOT_TICKED)
            break


def _set_sdt_text(sdt, text: str) -> None:
    """
    Replace the content of an SDT with plain text.

    Clears the SDT's <w:sdtContent> element and inserts a new paragraph
    with a single run containing the given text.  Preserves the <w:sdtPr>
    (properties) block so the field still looks like an SDT in Word.
    """
    content_tag = f"{{{W}}}sdtContent"
    p_tag       = f"{{{W}}}p"
    r_tag       = f"{{{W}}}r"
    t_tag       = f"{{{W}}}t"

    # Find or create <w:sdtContent>
    sdt_content = sdt.find(content_tag)
    if sdt_content is None:
        sdt_content = etree.SubElement(sdt, content_tag)

    # Remove existing children
    for child in list(sdt_content):
        sdt_content.remove(child)

    # Build <w:p><w:r><w:t>text</w:t></w:r></w:p>
    p = etree.SubElement(sdt_content, p_tag)
    r = etree.SubElement(p, r_tag)
    t = etree.SubElement(r, t_tag)
    t.text = text
    # Preserve whitespace
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")


def _set_table_cell_text(doc: Document, table_idx: int, row_idx: int, col_idx: int, text: str) -> None:
    """Replace all text in a table cell with the given string."""
    table = doc.tables[table_idx]
    cell  = table.rows[row_idx].cells[col_idx]

    # Clear all paragraphs except the first, then set text on first para
    for para in cell.paragraphs[1:]:
        p = para._element
        p.getparent().remove(p)

    first_para = cell.paragraphs[0]
    # Clear existing runs
    for run in first_para.runs:
        run.text = ""
    if first_para.runs:
        first_para.runs[0].text = text
    else:
        first_para.add_run(text)


# ── Public function ────────────────────────────────────────────────────────────

def generate_dpia_docx(plan: MappingPlan, stem: str) -> str:
    """
    Populate a fresh copy of the DPIA template and save it to outputs/.

    Args:
        plan: MappingPlan from dpia_mapper.build_mapping_plan()
        stem: original uploaded filename stem (used in output filename)

    Returns:
        Absolute path of the generated .docx file.
    """
    _OUTPUTS.mkdir(parents=True, exist_ok=True)

    today     = date.today().strftime("%Y-%m-%d")
    safe_stem = "".join(c if c.isalnum() or c in "-_" else "_" for c in stem)
    out_path  = _OUTPUTS / f"DPIA_{safe_stem}_{today}.docx"

    # Work on a fresh in-memory copy so we never touch the template file
    doc = Document(str(_TEMPLATE))

    # Collect SDTs once
    text_sdts     = _collect_text_sdts(doc)
    date_sdts     = _collect_date_sdts(doc)
    checkbox_sdts = _collect_checkbox_sdts(doc)

    logger.debug(
        "Found %d text-input SDTs, %d date-picker SDTs, %d checkbox SDTs",
        len(text_sdts), len(date_sdts), len(checkbox_sdts),
    )

    # ── Text-input SDT insertions ─────────────────────────────────────────────
    for sdt_idx, value in plan.sdt_insertions:
        if sdt_idx >= len(text_sdts):
            logger.warning("SDT index %d out of range (only %d text SDTs found) — skipping", sdt_idx, len(text_sdts))
            continue
        try:
            _set_sdt_text(text_sdts[sdt_idx], value)
            logger.debug("Set text SDT[%d] = %r", sdt_idx, value[:40])
        except Exception as exc:
            logger.warning("Failed to set text SDT[%d]: %s", sdt_idx, exc)

    # ── Date-picker SDT insertions ────────────────────────────────────────────
    for date_idx, value in plan.date_sdt_insertions:
        if date_idx >= len(date_sdts):
            logger.warning("Date SDT index %d out of range (only %d date SDTs found) — skipping", date_idx, len(date_sdts))
            continue
        try:
            _set_sdt_text(date_sdts[date_idx], value)
            logger.debug("Set date SDT[%d] = %r", date_idx, value[:40])
        except Exception as exc:
            logger.warning("Failed to set date SDT[%d]: %s", date_idx, exc)

    # ── Checkbox ticking ──────────────────────────────────────────────────────
    for cb_idx in plan.checkbox_ticks:
        if cb_idx >= len(checkbox_sdts):
            logger.warning("Checkbox index %d out of range (only %d checkbox SDTs found) — skipping", cb_idx, len(checkbox_sdts))
            continue
        try:
            _tick_checkbox(checkbox_sdts[cb_idx])
            logger.debug("Ticked checkbox SDT[%d]", cb_idx)
        except Exception as exc:
            logger.warning("Failed to tick checkbox SDT[%d]: %s", cb_idx, exc)

    # ── Table cell insertions ─────────────────────────────────────────────────
    for t_idx, r_idx, c_idx, value in plan.table_insertions:
        try:
            _set_table_cell_text(doc, t_idx, r_idx, c_idx, value)
            logger.debug("Set Table[%d][%d][%d] = %r", t_idx, r_idx, c_idx, value[:40])
        except Exception as exc:
            logger.warning("Failed to set Table[%d][%d][%d]: %s", t_idx, r_idx, c_idx, exc)

    doc.save(str(out_path))
    logger.info("Saved DPIA document to %s", out_path)
    return str(out_path)
